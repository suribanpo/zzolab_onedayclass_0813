import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PIL import Image

st.title("💼 명함 생성기 💼")

# 사용자 입력 받기
name = st.text_input("이름 또는 별명을 입력하세요:", placeholder="예: 홍길동 (홍홍)")
school_name = st.text_input("학교 이름을 입력하세요:", placeholder="예: 쪼랩고등학교")
email = st.text_input("메일주소를 입력하세요:", placeholder="예: example@korea.kr")
website = st.text_input("개인 홈페이지를 입력하세요:", placeholder="예: zzolab.com")
if st.button("🚀 명함 생성하기"):
    if name and school_name and email and website:
        with st.spinner("명함을 생성하는 중입니다..."):
            import time
            time.sleep(1)  # 2초 대기
            st.toast("명함이 생성되었어요!")  # 토스트 메시지
            # 워드 문서 생성
            doc = Document()

            # 문서 사이즈를 명함 크기로 설정 (9cm x 5cm)
            section = doc.sections[0]
            section.page_width = Cm(9)
            section.page_height = Cm(5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)

            # 명함 내용 추가
            content = doc.add_paragraph()
            content.add_run(f"이름: {name}\n근무지: {school_name}\nE-mail: {email}\nMy Page: {website}\n").bold = True
            content.add_run("코딩하는 선생님들 커뮤니티, 쪼랩 회원").italic = True
            content.runs[-1].font.size = Pt(6)  # 글씨 크기 조정
            content.runs[-1].font.color.rgb = RGBColor(128, 128, 128)  # 회색으로 설정
            # content.add_run("\n")  # 줄바꿈 추가
            content.runs[-1].font.size = Pt(8)  # 글씨 크기 조정
            content.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in content.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

            # 페이지 구분 추가
            doc.add_page_break()

            # 로고 이미지 삽입 (60% 사이즈로 정중앙 정렬)
            logo_path = "data/images/main_logo.png"  # 이미지 경로
            logo_img = Image.open(logo_path)
            logo_width_cm = Cm(4.8)  # 60% of max width (8cm)
            logo_height_cm = logo_width_cm * (logo_img.size[1] / logo_img.size[0])  # 비율 유지
            logo_para = doc.add_paragraph()
            logo_run = logo_para.add_run("\n\n")
            logo_run.add_picture(logo_path, width=logo_width_cm, height=logo_height_cm)
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph().add_run().add_break()  # 중앙 정렬을 위한 빈 줄 추가

            # 워드 문서를 메모리에 저장
            doc_file = BytesIO()
            doc.save(doc_file)
            doc_file.seek(0)

            # 워드 문서 다운로드 버튼
            st.download_button(
                label="📥 명함 워드파일 다운로드",
                data=doc_file,
                file_name="명함.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("❌ 모든 필드를 입력해주세요.")
else:
    st.error("❌ 모든 필드를 입력해주세요.")